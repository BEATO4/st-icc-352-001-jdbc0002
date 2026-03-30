#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Establece el color de fondo de una celda"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_style(doc, text, level=1):
    """Agrega un heading con estilo"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.style = 'Heading 1'
    elif level == 2:
        heading.style = 'Heading 2'
    return heading

def create_report():
    doc = Document()

    # ========================
    # ENCABEZADO
    # ========================

    # Tabla de encabezado
    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    table.allow_autofit = False

    # Fila 1: Estudiantes
    cell = table.cell(0, 0)
    cell.text = "Estudiantes:"
    cell_bold = cell.paragraphs[0].runs[0]
    cell_bold.bold = True

    cell = table.cell(0, 1)
    cell.text = "Rafael Ramírez - 10149810\nJostin Beato - 10150326"

    # Fila 2: Profesor
    cell = table.cell(1, 0)
    cell.text = "Profesor:"
    cell_bold = cell.paragraphs[0].runs[0]
    cell_bold.bold = True

    cell = table.cell(1, 1)
    cell.text = "Carlos Camacho"

    # Fila 3: Materia
    cell = table.cell(2, 0)
    cell.text = "Materia:"
    cell_bold = cell.paragraphs[0].runs[0]
    cell_bold.bold = True

    cell = table.cell(2, 1)
    cell.text = "Programación Web (ICC-352)"

    # Fila 4: Trabajo
    cell = table.cell(3, 0)
    cell.text = "Trabajo:"
    cell_bold = cell.paragraphs[0].runs[0]
    cell_bold.bold = True

    cell = table.cell(3, 1)
    cell.text = "Práctica #5 – AJAX + WebSocket"

    # Fila 5: Fecha
    cell = table.cell(4, 0)
    cell.text = "Fecha:"
    cell_bold = cell.paragraphs[0].runs[0]
    cell_bold.bold = True

    cell = table.cell(4, 1)
    cell.text = "30 de Marzo de 2026"

    doc.add_paragraph()  # Espacio

    # ========================
    # INTRODUCCIÓN
    # ========================

    doc.add_heading('Introducción', level=1)

    intro_text = """Esta práctica representa una extensión significativa de la Práctica 4, donde ya teníamos un sistema de blog completamente funcional con persistencia en base de datos. En esta ocasión, el objetivo principal es implementar un sistema de chat online en tiempo real utilizando WebSocket y mejorar la experiencia del usuario con Fetch API para operaciones asincrónicas AJAX sin necesidad de recargar la página.

El proyecto mantiene toda la funcionalidad anterior (autenticación, creación de artículos, comentarios, paginación, etiquetas) pero agrega una capa de comunicación en tiempo real que permite a los visitantes del blog chatear con administradores o autores sin necesidad de actualizar la página. Esto requirió arquitectura más sofisticada en el servidor, implementación de WebSocket, y una interfaz de usuario más compleja en el cliente.

Además se implementó un panel de administración para que los autores y administradores puedan gestionar múltiples conversaciones en paralelo, cerrar chats, y reabrirlos según sea necesario. Todo esto con persistencia completa en base de datos PostgreSQL cuando se ejecuta en Docker.

El proyecto sigue utilizando Javalin como framework web, Thymeleaf como motor de plantillas, Bootstrap 5.3.8 para el diseño visual, y ahora también incluye WebSocket nativo de Javalin para la comunicación en tiempo real."""

    doc.add_paragraph(intro_text)

    # ========================
    # DESARROLLO
    # ========================

    doc.add_heading('Desarrollo', level=1)

    # Sección: Estructura del Proyecto
    doc.add_heading('Estructura del Proyecto', level=2)

    structure_text = """La estructura de carpetas se extiende respecto a la práctica anterior para incluir nuevos componentes relacionados con el chat, WebSocket, y la integración con múltiples tecnologías:"""

    doc.add_paragraph(structure_text)

    # Código de estructura
    structure_code = """practica-5/
├── Dockerfile
├── docker-compose.yml
├── src/
│   └── main/
│       ├── java/org/pucmm/blog/
│       │   ├── encapsulaciones/
│       │   │   ├── Articulo.java
│       │   │   ├── ChatSesion.java (NUEVO)
│       │   │   ├── ChatMensaje.java (NUEVO)
│       │   │   ├── Comentario.java
│       │   │   ├── Etiqueta.java
│       │   │   └── Usuario.java
│       │   ├── servicios/
│       │   │   ├── ArticuloServices.java
│       │   │   ├── ChatSesionServices.java (NUEVO)
│       │   │   ├── ChatMensajeServices.java (NUEVO)
│       │   │   ├── AuditoriaServices.java
│       │   │   ├── GestionDb.java
│       │   │   └── UsuarioServices.java
│       │   └── Main.java
│       └── resources/
│           ├── META-INF/
│           │   └── persistence.xml
│           ├── public/
│           │   ├── css/
│           │   └── js/
│           │       ├── chat-widget.js (NUEVO)
│           │       ├── admin-chat.js (NUEVO)
│           │       ├── index-pagination.js
│           │       └── si.css
│           └── templates/
│               ├── admin_chats.html (NUEVO)
│               ├── articulo.html
│               ├── crear_articulo.html
│               ├── crear_usuario.html
│               ├── index.html
│               ├── layout.html
│               └── login.html
├── .gitignore
├── build.gradle.kts
├── gradlew
├── gradlew.bat
└── README.md"""

    p = doc.add_paragraph(structure_code)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Sección: Tecnologías utilizadas
    doc.add_heading('Tecnologías Utilizadas', level=2)

    # Tabla de tecnologías
    tech_table = doc.add_table(rows=10, cols=2)
    tech_table.style = 'Light Grid Accent 1'

    tech_table.cell(0, 0).text = "Tecnología"
    tech_table.cell(0, 1).text = "Descripción"

    techs = [
        ("Javalin 6.7.0", "Framework web ligero para Java. Usado para definir rutas HTTP, manejar WebSocket y procesar requests."),
        ("WebSocket (Javalin)", "Implementación de WebSocket nativa en Javalin para comunicación bidireccional en tiempo real."),
        ("Thymeleaf 3.1.2", "Motor de plantillas para renderizar vistas dinámicas con datos del servidor."),
        ("Fetch API (JavaScript)", "API moderna del navegador para hacer peticiones HTTP asincrónicas (AJAX) sin recargar página."),
        ("Hibernate 6 / JPA 3", "ORM que mapea entidades Java a tablas en base de datos."),
        ("PostgreSQL 16", "Base de datos relacional en Docker Compose para persistencia en producción."),
        ("H2 Database", "Base de datos embebida para ejecución local sin dependencias externas."),
        ("Docker + Docker Compose", "Contenedorización y orquestación de servicios (app + PostgreSQL)."),
        ("Bootstrap 5.3.8", "Framework CSS para diseño visual responsivo y consistente."),
    ]

    for i, (tech, desc) in enumerate(techs, 1):
        tech_table.cell(i, 0).text = tech
        tech_table.cell(i, 1).text = desc

    doc.add_paragraph()  # Espacio

    # Sección: Funcionalidades Implementadas
    doc.add_heading('Funcionalidades Implementadas', level=2)

    # 1. AJAX Paginación
    doc.add_heading('1. Paginación AJAX sin recargar página', level=3)

    ajax_text = """El endpoint GET /api/articulos?page=N retorna un JSON con los artículos de la página solicitada. El cliente usa Fetch API para obtener los datos y actualiza el DOM dinámicamente sin recargar la página completa. Esto mejora significativamente la experiencia del usuario al permitir navegar entre páginas de artículos de forma fluida e instantánea."""

    doc.add_paragraph(ajax_text)

    # 2. Chat en Tiempo Real
    doc.add_heading('2. Sistema de Chat en Tiempo Real con WebSocket', level=3)

    chat_text = """Se implementó un widget de chat que aparece tanto en la página principal como en la visualización de cada artículo. Los visitantes pueden ingresar su nombre y comenzar una conversación con administradores o autores. La comunicación se realiza mediante WebSocket (/ws/chat) que mantiene una conexión bidireccional abierta para enviar y recibir mensajes en tiempo real sin latencia.

Características:
• Cualquier visitante (autenticado o no) puede usar el chat
• Cada sesión de chat tiene un token único para seguridad
• Mensajes se persisten en base de datos
• Soporta múltiples chats simultáneos
• Reconexión automática si se pierde la conexión"""

    doc.add_paragraph(chat_text)

    # 3. Panel Admin
    doc.add_heading('3. Panel de Administración de Chats', level=3)

    admin_text = """Accesible en /admin/chats (solo para admin/autor), este panel muestra todas las conversaciones activas y recientes. El admin puede:
• Ver lista de chats abiertos
• Ver historial completo de mensajes de cada chat
• Responder en tiempo real a múltiples visitantes
• Cerrar un chat (visitante ve mensaje de cierre)
• Reabrir un chat cerrado
• Eliminar un chat (solo si está cerrado)

El panel actualiza automáticamente la lista de chats sin recargar."""

    doc.add_paragraph(admin_text)

    # 4. Control de Permisos
    doc.add_heading('4. Control de Permisos Basado en Roles', level=3)

    perms_text = """Se implementó un sistema de permisos en el backend que valida:
• Visitantes: Pueden enviar mensajes, pero no crear/cerrar/eliminar chats
• Admin/Autor: Pueden crear nuevos chats, responder, cerrar y reabrir
• Solo Admin: Puede eliminar chats

En el frontend, los botones de administración se ocultan si el usuario no tiene permisos."""

    doc.add_paragraph(perms_text)

    # 5. Docker
    doc.add_heading('5. Contenedorización con Docker Compose', level=3)

    docker_text = """El proyecto incluye Dockerfile con build multistage (compilación con Gradle + ejecución con JRE) y docker-compose.yml que orquesta dos servicios:
• app: Aplicación Javalin en puerto 7000
• db: PostgreSQL 16-alpine en puerto 5432

La base de datos persiste en un volumen Docker, permitiendo datos consistentes entre reinicios."""

    doc.add_paragraph(docker_text)

    # ========================
    # PROBLEMAS ENCONTRADOS
    # ========================

    doc.add_heading('Problemas Encontrados y Soluciones', level=2)

    problems = [
        ("Sincronización de sesiones WebSocket", "Las sesiones de WebSocket se perdían entre reconexiones. Se implementó validación de tokens en BD y reconexión automática del cliente."),
        ("Permisos no se validaban en cliente", "Visitantes podían ver botones de admin aunque no tuvieran permisos. Se agregó detección de rol en frontend para ocultar botones."),
        ("Llamadas AJAX fallaban", "El cliente enviaba formato de datos incorrecto. Se corrigió el payload JSON para incluir todos los campos requeridos."),
        ("Input deshabilitado siempre", "El input de mensajes nunca se habilitaba. Se agregó lógica en actualización de UI para habilitar/deshabilitar según estado de sesión."),
        ("JDBC_DATABASE_URL no configurada en Docker", "AuditoriaServices fallaba en contenedor. Se agregó la variable de entorno en docker-compose.yml."),
    ]

    for prob, sol in problems:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(prob + ": ").bold = True
        p.add_run(sol)

    # ========================
    # CONCLUSIÓN
    # ========================

    doc.add_heading('Conclusión', level=2)

    conclusion_text = """Esta práctica representó un salto importante en complejidad respecto a la Práctica 4. Pasar de un blog estático con CRUD básico a un sistema con comunicación bidireccional en tiempo real requirió entendimiento profundo de:

1. WebSocket: Protocolo de comunicación de baja latencia para aplicaciones real-time
2. Asincronía: JavaScript promises, Fetch API, y manejo de callbacks
3. Arquitectura: Cómo separar responsabilidades entre frontend (interacción) y backend (lógica de negocio)
4. Seguridad: Validación de permisos, tokens únicos por sesión, control de acceso
5. Infraestructura: Docker Compose para orquestación multi-servicio
6. Debugging: Uso de DevTools de navegador para ver WebSocket connections

La implementación de múltiples chats simultáneos fue especialmente interesante, ya que requirió estructuras de datos especializadas en el servidor (Sets y Maps concurrentes) y sincronización cuidadosa entre clientes.

El proyecto ahora es mucho más cercano a aplicaciones reales en producción, con persistencia, seguridad, comunicación en tiempo real, y buenas prácticas de desarrollo. El uso de Docker Compose permite que el proyecto sea reproducible en cualquier máquina sin dependencias de instalación local.

Esta práctica consolidó conceptos fundamentales de desarrollo web moderno y sentó las bases para trabajar con tecnologías más avanzadas en el futuro."""

    doc.add_paragraph(conclusion_text)

    # ========================
    # REPOSITORIO Y BIBLIOGRAFÍA
    # ========================

    doc.add_heading('Link del Repositorio de GitHub', level=2)
    p = doc.add_paragraph()
    p.add_run("https://github.com/BEATO4/st-icc-352-001-jdbc0002.git").bold = True

    doc.add_heading('Bibliografía', level=2)

    bibliography = [
        "Javalin. (2024). Javalin – A simple web framework for Java and Kotlin. https://javalin.io/",
        "Thymeleaf. (2024). Thymeleaf Documentation. https://www.thymeleaf.org/documentation.html",
        "Bootstrap. (2024). Bootstrap 5.3 Documentation. https://getbootstrap.com/docs/5.3/",
        "MDN Web Docs. (2024). WebSocket API. https://developer.mozilla.org/en-US/docs/Web/API/WebSocket",
        "MDN Web Docs. (2024). Fetch API. https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API",
        "Hibernate. (2024). Hibernate ORM Documentation. https://hibernate.org/orm/documentation/",
        "PostgreSQL. (2024). PostgreSQL Documentation. https://www.postgresql.org/docs/",
        "Docker. (2024). Docker Documentation. https://docs.docker.com/",
        "Gradle Inc. (2024). Gradle Build Tool. https://docs.gradle.org/current/userguide/userguide.html",
        "Camacho, C. (2026). Material de la Asignatura ICC-352 Programación Web. PUCMM – Facultad de Ciencias e Ingeniería.",
    ]

    for item in bibliography:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item)

    # Guardar documento
    output_path = "C:\\Users\\rafav\\IdeaProjects\\st-icc-352-001-jdbc0002\\practica-5\\Practica_5_AJAX_WebSocket.docx"
    doc.save(output_path)
    print(f"✅ Documento creado exitosamente: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()

